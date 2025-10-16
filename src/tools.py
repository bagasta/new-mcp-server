"""Calculator and web utility tools for the MCP server."""

from __future__ import annotations

import base64
import math
import os
import re
import unicodedata
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Union
from urllib.parse import urlparse

import httpx
from docx import Document
from fpdf import FPDF
from fpdf.errors import FPDFException


Number = Union[int, float]

_ZERO_WIDTH_CODEPOINTS = {
    0x200B,  # ZERO WIDTH SPACE
    0x200C,  # ZERO WIDTH NON-JOINER
    0x200D,  # ZERO WIDTH JOINER
    0x2060,  # WORD JOINER
    0xFEFF,  # ZERO WIDTH NO-BREAK SPACE
}
_VARIATION_RANGES: tuple[tuple[int, int], ...] = (
    (0xFE00, 0xFE0F),
    (0xE0100, 0xE01EF),
)


class Calculator:
    """Calculator with basic and advanced mathematical operations."""

    @staticmethod
    def add(a: Number, b: Number) -> float:
        """Add two numbers.

        Args:
            a: The first addend.
            b: The second addend.

        Returns:
            The sum of ``a`` and ``b``.
        """
        return float(a + b)

    @staticmethod
    def subtract(a: Number, b: Number) -> float:
        """Subtract one number from another.

        Args:
            a: The value to subtract from.
            b: The value to subtract.

        Returns:
            The result of ``a - b``.
        """
        return float(a - b)

    @staticmethod
    def multiply(a: Number, b: Number) -> float:
        """Multiply two numbers.

        Args:
            a: The first factor.
            b: The second factor.

        Returns:
            The product of ``a`` and ``b``.
        """
        return float(a * b)

    @staticmethod
    def divide(a: Number, b: Number) -> float:
        """Divide one number by another.

        Args:
            a: The dividend.
            b: The divisor.

        Returns:
            The quotient of ``a / b``.

        Raises:
            ValueError: If ``b`` is zero.
        """
        if b == 0:
            raise ValueError("Cannot divide by zero.")
        return float(a / b)

    @staticmethod
    def power(base: Number, exponent: Number) -> float:
        """Raise a base to a power.

        Args:
            base: The base value.
            exponent: The exponent value.

        Returns:
            ``base`` raised to ``exponent``.
        """
        return float(math.pow(base, exponent))

    @staticmethod
    def sqrt(value: Number) -> float:
        """Calculate the square root of a non-negative number.

        Args:
            value: The value whose square root should be calculated.

        Returns:
            The square root of ``value``.

        Raises:
            ValueError: If ``value`` is negative.
        """
        if value < 0:
            raise ValueError("Cannot calculate square root of a negative number.")
        return math.sqrt(value)

    @staticmethod
    def factorial(value: int) -> int:
        """Calculate the factorial of a non-negative integer.

        Args:
            value: The integer whose factorial should be calculated.

        Returns:
            The factorial of ``value``.

        Raises:
            ValueError: If ``value`` is negative.
        """
        if value < 0:
            raise ValueError("Factorial is undefined for negative values.")
        return math.factorial(value)

    @staticmethod
    def percentage(value: Number, percent: Number) -> float:
        """Calculate the percentage of a value.

        Args:
            value: The base value.
            percent: The percentage to apply.

        Returns:
            The result of ``percent`` percent of ``value``.
        """
        return float((value * percent) / 100)


class WebFetcher:
    """Utility for retrieving web content over HTTP."""

    @staticmethod
    async def fetch(url: str, timeout: float = 10.0) -> dict[str, str | int]:
        """Fetch textual content from a URL.

        Args:
            url: Fully-qualified HTTP or HTTPS URL to retrieve.
            timeout: Maximum time to wait for the response, in seconds.

        Returns:
            A dictionary containing the final URL, HTTP status code, and response body.

        Raises:
            ValueError: If the URL scheme is unsupported or the request fails.
        """
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"}:
            raise ValueError("URL must use http or https.")

        try:
            async with httpx.AsyncClient(follow_redirects=True, timeout=timeout) as client:
                response = await client.get(url)
                response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            raise ValueError(
                f"Request to {url} failed with status code {exc.response.status_code}."
            ) from exc
        except httpx.HTTPError as exc:
            raise ValueError(f"Request to {url} failed: {exc}") from exc

        max_chars = 500
        content = response.text
        truncated = len(content) > max_chars
        if truncated:
            content = content[:max_chars]

        note = (
            f"Content truncated to {max_chars} characters for processing."
            if truncated
            else "Content retrieved successfully."
        )

        return {
            "url": str(response.url),
            "status_code": response.status_code,
            "content_snippet": content,
            "note": note,
        }


class WebSearcher:
    """Utility for performing Serper-powered web searches."""

    _DEFAULT_ENDPOINT = "https://google.serper.dev/search"

    def __init__(self, api_key: str | None = None, endpoint: str | None = None) -> None:
        self._api_key = api_key
        self._endpoint = endpoint or self._DEFAULT_ENDPOINT

    def _resolve_api_key(self) -> str:
        api_key = self._api_key or os.getenv("SERPER_API_KEY")
        if not api_key:
            raise ValueError(
                "SERPER_API_KEY is not configured. "
                "Set the environment variable or provide an API key when constructing WebSearcher."
            )
        return api_key

    async def search(
        self,
        query: str,
        *,
        country: str = "us",
        language: str = "en",
        num_results: int = 5,
        timeout: float = 10.0,
    ) -> dict[str, Any]:
        """Execute a web search using the Serper API."""
        if not query.strip():
            raise ValueError("Query must not be empty.")

        limit = max(1, min(num_results, 10))
        payload = {
            "q": query,
            "gl": country.lower(),
            "hl": language.lower(),
            "num": limit,
        }

        headers = {
            "X-API-KEY": self._resolve_api_key(),
            "Content-Type": "application/json",
        }

        timeout_config = httpx.Timeout(timeout, connect=timeout)

        try:
            async with httpx.AsyncClient(timeout=timeout_config) as client:
                response = await client.post(self._endpoint, json=payload, headers=headers)
                response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            raise ValueError(
                f"Serper API request failed with status {exc.response.status_code}: {exc.response.text}"
            ) from exc
        except httpx.TimeoutException as exc:
            raise ValueError(
                "Serper API request timed out. Check your network connectivity and API key configuration."
            ) from exc
        except httpx.HTTPError as exc:
            raise ValueError(f"Serper API request failed: {exc}") from exc

        data = response.json()
        organic_results = data.get("organic", []) or []
        selected_results = organic_results[:limit]

        results = [
            {
                "title": item.get("title"),
                "link": item.get("link"),
                "snippet": item.get("snippet"),
                "source": item.get("source"),
            }
            for item in selected_results
        ]

        answer_box = data.get("answerBox") or {}
        knowledge_graph = data.get("knowledgeGraph") or {}

        return {
            "query": query,
            "parameters": {
                "country": country,
                "language": language,
                "requested_results": num_results,
                "returned_results": len(results),
            },
            "results": results,
            "answer_box": answer_box,
            "knowledge_graph": knowledge_graph,
        }


class PDFGenerator:
    """Utility for generating simple PDF documents."""

    _FONT_PATH = Path("assets/fonts/NotoSans-Regular.ttf")
    _FONT_NAME = "NotoSans"
    _MAX_WORD_CHUNK = 80

    def __init__(self, default_author: str = "MCP Calculator Server") -> None:
        self._default_author = default_author
        self._ensure_font_available()
        self._replaced_glyphs = False

    def _ensure_font_available(self) -> None:
        if not self._FONT_PATH.exists():
            raise FileNotFoundError(
                f"PDF font file not found at {self._FONT_PATH}. "
                "Ensure the required font is bundled with the application."
            )

    def generate_pdf(
        self,
        *,
        title: str,
        content: str,
        filename: str | None = None,
        author: str | None = None,
    ) -> dict[str, Any]:
        """Generate a PDF document containing the provided content."""
        normalized_title = title.strip()
        if not normalized_title:
            raise ValueError("Title must not be empty.")

        if not content.strip():
            raise ValueError("Content must not be empty.")

        self._replaced_glyphs = False

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.add_font(self._FONT_NAME, style="", fname=str(self._FONT_PATH), uni=True)
        pdf.add_font(self._FONT_NAME, style="B", fname=str(self._FONT_PATH), uni=True)
        pdf.set_font(self._FONT_NAME, "B", 16)
        pdf.set_title(normalized_title)
        pdf.set_author(author or self._default_author)

        self._safe_multicell(pdf, normalized_title, 10)
        pdf.ln(5)

        pdf.set_font(self._FONT_NAME, size=12)
        for paragraph in content.split("\n\n"):
            lines = paragraph.splitlines() or [paragraph]
            for line in lines:
                self._safe_multicell(pdf, line, 8)
            pdf.ln(4)

        raw_pdf = pdf.output(dest="S")
        pdf_bytes = raw_pdf.encode("latin1") if isinstance(raw_pdf, str) else raw_pdf

        safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", normalized_title.lower()).strip("_") or "document"
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        final_filename = filename or f"{safe_title}_{timestamp}.pdf"

        encoded = base64.b64encode(pdf_bytes).decode("ascii")

        note = "PDF generated successfully."
        if self._replaced_glyphs:
            note += " Some unsupported characters were omitted."

        return {
            "filename": final_filename,
            "mime_type": "application/pdf",
            "base64_content": encoded,
            "size_bytes": len(pdf_bytes),
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "note": note,
        }

    def _force_wrap(self, text: str) -> str:
        if not text:
            return text

        tokens = re.split(r"(\s+)", text)
        wrapped_tokens: list[str] = []
        for token in tokens:
            if not token or token.isspace() or len(token) <= self._MAX_WORD_CHUNK:
                wrapped_tokens.append(token)
                continue
            chunks = [
                token[i : i + self._MAX_WORD_CHUNK]
                for i in range(0, len(token), self._MAX_WORD_CHUNK)
            ]
            wrapped_tokens.append(" ".join(chunks))
        return "".join(wrapped_tokens)

    def _safe_multicell(self, pdf: FPDF, text: str, line_height: float) -> None:
        sanitized = self._sanitize_text(pdf, text)
        try:
            pdf.multi_cell(0, line_height, sanitized)
            return
        except FPDFException:
            pass

        forced = self._force_wrap(sanitized)
        if forced != sanitized:
            try:
                pdf.multi_cell(0, line_height, forced)
                return
            except FPDFException:
                pass

        # Final fallback: break text based on rendered width.
        for chunk in self._chunk_by_width(pdf, sanitized):
            try:
                pdf.multi_cell(0, line_height, chunk)
            except FPDFException:
                self._fallback_ascii(pdf, chunk, line_height)

    def _chunk_by_width(self, pdf: FPDF, text: str) -> list[str]:
        if not text:
            return [text]

        max_width = pdf.w - pdf.r_margin - pdf.l_margin
        if max_width <= 0:
            return [text]

        chunks: list[str] = []
        current = ""
        current_width = 0.0

        for char in text:
            char_width = pdf.get_string_width(char)
            if char_width <= 0:
                # Fallback to approximate width using a placeholder character.
                char_width = pdf.get_string_width(" ")
                if char_width <= 0:
                    char_width = 1.0

            if current and current_width + char_width > max_width:
                chunks.append(current)
                current = char
                current_width = char_width
            else:
                current += char
                current_width += char_width

        if current:
            chunks.append(current)

        return chunks

    def _sanitize_text(self, pdf: FPDF, text: str) -> str:
        if not text:
            return text

        sanitized_chars: list[str] = []
        max_width = pdf.w - pdf.r_margin - pdf.l_margin
        if max_width <= 0:
            max_width = 1.0
        space_width = pdf.get_string_width(" ") or 1.0

        for char in text:
            codepoint = ord(char)
            if any(start <= codepoint <= end for start, end in _VARIATION_RANGES):
                self._replaced_glyphs = True
                continue
            if codepoint in _ZERO_WIDTH_CODEPOINTS:
                self._replaced_glyphs = True
                continue

            category = unicodedata.category(char)
            if category in {"Cf", "Cc"} and char not in {"\n", "\r", "\t"}:
                self._replaced_glyphs = True
                continue

            width = pdf.get_string_width(char)
            if width <= 0:
                self._replaced_glyphs = True
                # Skip zero-width or unsupported glyphs entirely.
                continue

            if width > max_width:  # excessively wide glyph, replace with placeholder
                placeholder = "?"
                placeholder_width = pdf.get_string_width(placeholder) or space_width
                sanitized_chars.append(placeholder)
                self._replaced_glyphs = True
                continue

            sanitized_chars.append(char)

        # If everything was stripped, return a single placeholder to keep layout.
        if not sanitized_chars:
            self._replaced_glyphs = True
            return "?"

        return "".join(sanitized_chars)

    def _fallback_ascii(self, pdf: FPDF, text: str, line_height: float) -> None:
        """Last-resort writer that strips unsupported glyphs."""
        ascii_text = text.encode("ascii", "replace").decode("ascii")
        if ascii_text != text:
            self._replaced_glyphs = True

        try:
            pdf.multi_cell(0, line_height, ascii_text)
            return
        except FPDFException:
            self._replaced_glyphs = True

        # Final final fallback: write character-by-character.
        x_start = pdf.get_x()
        max_width = pdf.w - pdf.r_margin - pdf.l_margin
        for char in ascii_text or "?":
            if char == "\n":
                pdf.ln(line_height)
                pdf.set_x(x_start)
                continue

            width = pdf.get_string_width(char) or 1.0
            if pdf.get_x() + width > pdf.w - pdf.r_margin:
                pdf.ln(line_height)
                pdf.set_x(x_start)

            pdf.cell(width, line_height, char, ln=0)
        pdf.ln(line_height)


class DOCXGenerator:
    """Utility for generating simple DOCX documents."""

    def __init__(self, default_author: str = "MCP Calculator Server") -> None:
        self._default_author = default_author

    def generate_docx(
        self,
        *,
        title: str,
        content: str,
        filename: str | None = None,
        author: str | None = None,
    ) -> dict[str, Any]:
        """Generate a DOCX document containing the provided content."""
        normalized_title = title.strip()
        if not normalized_title:
            raise ValueError("Title must not be empty.")

        if not content.strip():
            raise ValueError("Content must not be empty.")

        document = Document()
        properties = document.core_properties
        properties.author = author or self._default_author
        properties.title = normalized_title

        document.add_heading(normalized_title, level=0)

        for block in content.split("\n\n"):
            lines = block.splitlines()
            if not lines:
                document.add_paragraph("")
                continue

            paragraph = document.add_paragraph()
            for index, line in enumerate(lines):
                if index:
                    paragraph.add_run().add_break()
                paragraph.add_run(line)

        buffer = BytesIO()
        document.save(buffer)
        docx_bytes = buffer.getvalue()

        safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", normalized_title.lower()).strip("_") or "document"
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        final_filename = filename or f"{safe_title}_{timestamp}.docx"

        encoded = base64.b64encode(docx_bytes).decode("ascii")

        return {
            "filename": final_filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "base64_content": encoded,
            "size_bytes": len(docx_bytes),
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "note": "DOCX generated successfully.",
        }
